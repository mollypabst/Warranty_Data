import datetime

import requests
import json
import base64
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import matplotlib.dates as mdate
import time
import matplotlib.ticker as mtick
from operator import itemgetter 
import os

#  API_ADDRESS = "https://api.smartsensor.abb.com"
from docx import *
from docx.shared import Inches



def authenticate(api_key, addr):
    # Returns authentication token for the given API Key
    url = f"{addr}/Auth/Key"
    body = {'apiKey': api_key, 'deviceUID': "string"}
    # Making a request
    r = requests.post(url, json=body)
    if r.status_code == 200:
        # Converting json response to dictionary
        result = json.loads(r.text)
        # Getting the token from the dictionary
        token = result["authToken"]
        return token
    else:
        return ""


def collect_latest_raw_data(api_key, addr, sn, id):
    token = authenticate(api_key, addr)
    url = f"{addr}/Sensor/Feature/Value/{sn}"
    params = {'sensorTypeID': id, 'featureTypes': "RawDataCollectionCSV", 'complexObject': True}
    header = {'Authorization': f'Bearer {token}'}

    r = requests.get(url, params=params, headers=header)
    if r.status_code == 200:
        result = json.loads(r.text)
        # Getting the file name and content from the response
        file_name = result[0]["featureValue"][0]['fileName']
        content_encoded = result[0]["featureValue"][0]['fileContent']
        # Decoding base64 content
        content_bytes = content_encoded.encode('ascii')
        file_bytes = base64.b64decode(content_bytes)
        file_content = file_bytes.decode('ascii')
        # Writing csv file
        with open(file_name, 'w') as f:
            f.write(file_content)
        print(f"File {file_name} successfully written")


def get_sensor_feature(api_key, addr, sn, id, feature_types):
    token = authenticate(api_key, addr)
    url = f"{addr}/Sensor/Feature/{sn}"
    params = {'sensorTypeID': id, 'featureTypes': "SensorSelfTest", }
    header = {'Authorization': f'Bearer {token}'}

    response = requests.get(url, params=params, headers=header)

    if response.content is b'':
        return None
    response_json = json.loads(response.content)

    # Return None if the response code indicates an error
    if response.status_code != 200:
        # self._logger.debug('Error: Response Code ' + str(response.status_code) + " " + response_json)
        return None
    return response_json

    # return requests.get('Sensor/Feature/' + str(sn), {'sensorTypeID': id,  #                                                                 'featureTypes': feature_types})


def collect_all_raw_data(api_key, addr , sn, id, start_date, end_date):
    token = authenticate(api_key, addr)
    url = f"{adde}/Sensor/Feature/{sn}"
    params = {'sensorTypeID': id, 'featureTypes': "RawDataCollectionCSV", 'from': start_date, 'to': end_date,
        'complexObject': True}
    header = {'Authorization': f'Bearer {token}'}

    r = requests.get(url, params=params, headers=header)
    if r.status_code == 200:
        result = json.loads(r.text)
        # Getting the file name and content from the response
        for res in result:
            file_name = res["featureValue"][0]['fileName']
            content_encoded = res["featureValue"][0]['fileContent']
            # Decoding base64 content
            content_bytes = content_encoded.encode('ascii')
            file_bytes = base64.b64decode(content_bytes)
            file_content = file_bytes.decode('ascii')
            # Writing csv file
            with open(file_name, 'w') as f:
                f.write(file_content)
            print(f"File {file_name} successfully written")


def flatten_json(y):
    out = {}

    def flatten(x, name=''):
        if type(x) is dict:
            for a in x:
                flatten(x[a], name + a + '_')
        elif type(x) is list:
            i = 0
            for a in x:
                flatten(a, name + str(i) + '_')
                i += 1
        else:
            out[name[:-1]] = x

    flatten(y)
    return out


def main():
    # User input:
    #   api_key - key obtained from portal,
    #   sensor_serial_number - sensor that
    #                          we want to read raw data from
    #   sensor_type_id - the id of the sensor type.
    #                    For mounted bearigs the type is 3
    #prod_key = "TXpabU5HTXhNV010WTJFek1pMDBZMkZtTFdJeFpEQXROekUyWkRJNFlUSTNPRFZq"
    prod_key = "WkdKaE9XVmtZemt0WTJSbE5pMDBObVk0TFdJd09UTXRObUl3WWpjMVpETXlORE5s"
    beta_key = "TjJFek1UUm1ZamN0TkdGa01TMDBaamxtTFRrNVpEQXRNamN5WVdFeFpXSm1OekU1"

    sensor_list = [# Back test stand sensors from read_warr.py
    # ['10006245'],
    # ['10002402'],
    # ['10002617'],
    # ['10012626'],
    # ['10001880'],
    # ['10000275'],
    # ['10004413'],
    # ['10000156'],
    # ['10001188'],
    # ['10003563'],
    # ['10009399'],
    # ['10000705'],
    # ['10000155'],
    # ['10001991'],
    # ['10001217'],
    # ['10010020'],
    # ['10002635'],
    # ['10009991'],
    # ['10003992'],
    # ['10001206'],
    # ['10002879'],
    # ['10001219'],
    # ['10002631'],
    # ['10002602'],
    # ['10001975'],
    # ['10001216'],
    # ['10000709'],
    # ['10002620'],
    # ['10001213'],
    # ['10001773'],
    # ['10002633'],
    # ['10000704'],
    # ['10006524'],
    # ['10001211'],
    # ['10010018'],
    # ['10001979'],
    # ['10004938'],
    # ['10001222'],
    # ['10024969'],
    # ['10000887'],
    # ['10002628'],
    # ['10004992'],
    # ['10001218'],
    # ['10014754'],
    # ['10000710'],
    # ['10003697'],
    # ['10002896'],
    # ['10014836'],
    # ['10001891'],
    # ['10002630'],
    # ['10004982'],
    # ['10004130'],
    # ['10002247'],

    # ['10001798'],  # GSI bucket elevatior
    # ['10003280'], # prod app test with 200k connections

        

    # ['10002374'], # x1 (prod)
    # ['10003551'], # X2 (beta)
    # ['10003563'],  # dead
    # ['10003562'],
    # ['10002372'],
    # ['10003560'],
    # ['10003559'],
    # ['10003558'],
#     ['10002353'],  # dead
#     ['10003556'],  # dead
#     ['10003555'],
#     ['10003554'],
#     ['10002351'],
#     ['10003552'],
#     ['10002346'],
#     ['10003550'],
#     ['10002349'],
#     ['10003548'],
#     ['10003547'],
#     ['10001792'],
#     ['10002373'],
#     ['10003581'],
#     ['10003580'],
#     ['10003579'],
#     ['10002354'],
#     ['10003577'],
#     ['10003576'],
#     ['10002336'],
#     ['10002352'],
#     ['10003573'],
#     ['10003572'],
#     ['10003571'],
#     ['10002350'],
#     ['10002347'],
#     ['10003568'],
#     ['10003567'],
#     ['10002348'],
#     ['10003569'],
#     ['10002339'],
#     ['10001793'],
#     ['10002340'],
#     ['10003564'],
#     ['10003561'],
#     ['10003578'],
#     ['10002345'],
#     ['10003549'],
#     ['10003553'],
#     ['10003570'],
#     ['10002344'],
#     ['10003566'],
#     ['10003565'],
#     ['10003582'],
#     ['10003557'],
#     ['10003574'],
#     ['10002343'],
#     ['10002341'],
#     ['10003575'],
#     ['10002335'],
#     ['10003546'],
#     ['10001797'],

#    # Penn waste
#     ['10004471'],
#     ['10004472'],
#     ['10004473'],
#     ['10004474'],
#     ['10004475'],
#     ['10004476'],
#     ['10004477'],
#     ['10004478'],
#     ['10004479'],
#     ['10004480'],
#     ['10004481'],
#     ['10004482'],
#     ['10004483'],
#     ['10004484'],
#     ['10004485'],
#     ['10004486'],
#     ['10004487'],
#     ['10004488'],
#     ['10004489'],
#     ['10004490'],
#     ['10004509'],
#     ['10004511'],
#     ['10004512'],
#     ['10004513'],
#     ['10004514'],
#     ['10004515'],
#     ['10004516'],
#     ['10004517'],
#     ['10004518'],
#     ['10004519'],
#     ['10004520'],
#     ['10004521'],
#     ['10004522'],
#     ['10004523'],
#     ['10004524'],
#     ['10004525'],
#     ['10004526'],
#     ['10004527'],
#     ['10004528'],
#     ['10004529'],
#     ['10004530'],
#     ['10004531'],
#     ['10004532'],
#     ['10004533'],
#     ['10004534'],
#     ['10004535'],
#     ['10004536'],
#     ['10004537'],
#     ['10004538'],
#     ['10004539'],
#     ['10004540'],
#     ['10004541'],
#     ['10004542'],
#     ['10004543'],
#     ['10004544'],
#     ['10004545'],
#     ['10004546'],
#     ['10004547'],
#     ['10004548'],
#     ['10004549'],
#     ['10004550'],
#     ['10008093'],
#     ['10009395'],
#     ['10015941'],
#     ['10015942'],
#     ['10021469'],

#    # Rogers group
#     ['10004847'],
#     ['10004848'],
#     ['10004849'],
#     ['10004852'],
#     ['10004853'],
#     ['10004854'],
#     ['10004855'],
#     ['10004856'],
#     ['10004857'],
#     ['10004858'],
#     ['10004859'],
#     ['10004860'],
#     ['10004861'],
#     ['10004866'],
#     ['10004867'],
#     ['10004868'],
#     ['10004870'],
#     ['10004871'],
#     ['10004873'],
#     ['10004874'],
#     ['10004875'],
#     ['10004877'],
#     ['10004881'],
#     ['10004883'],
#     ['10004885'],
#     ['10009368'],
#     ['10009370'],
#     ['10009391'],
#     ['10009394'],
#     ['10009399'],
#     ['10015935'],
    # ['10015938'],
    # ['10021723'],
    # ['10021725'],

    # #DGV plant
    # ['50:31:AD:02:06:45', '10001084'],
    # ['10002879'],
    # ['10002880'],
    ['10002881'],
    ['10002882'],
    ['10002883'],
    ['10002886'],
#     ['10002887'],
#     ['10002891'],
#     ['10002892'],
#     ['10002896'],
#     ['10009367'],
#     ['10021823'],
#     ['10021824'],
#     ['10021826'],
    # ['10021837'],
    # ['10021840'],
    # ['10021841'],
    # ['10021849'],
    # ['10021852'],
    # ['10021854'],
    # ['10021855'],
    # ['10021856'],
    # ['10004056'], 
    ['10004055'], # 260J Centrifuge non drive
    ['10004051'],  # 320 centerfuge non drive side
    ['10004052'],  # 320 centerfuge. Drive side

    ]

    # delete old docs
    if os.path.exists("./all_sensors.docx"):
        os.remove("./all_sensors.docx")
    if os.path.exists("./dead_sensors.docx"):
        os.remove("./dead_sensors.docx")
    if os.path.exists("./alive_sensors.docx"):
        os.remove("./alive_sensors.docx")

    cnt = 0
    doc = Document()

    for sn in sensor_list:

        if cnt % 2:
            api_add = "https://beta.api.smartsensor.abb.com"
            api_key = beta_key
        else:
            api_add = "https://api.smartsensor.abb.com"
            api_key = prod_key

        api_add = "https://api.smartsensor.abb.com"
        api_key = prod_key
        print(api_add)
        print(api_key)


        sensor_serial_number = sn[0]
        sensor_type_id = 3
        start_date = "2020.01.13"
        end_date = "2020.09.23"
        feature_values = get_sensor_feature(api_key, api_add, sensor_serial_number, 3, 'SensorSelfTest')

        #print(feature_values)
        f_sort = sorted(feature_values, key=itemgetter('featureKeySequenceNo'))
        #print(f_sort)


        lst = []
        # csv = open("test_csv.csv", "w")  # "w" indicates that you're writing strings to the file
        for f in f_sort:
            lst.append(flatten_json(f['featureValue']))

        lst_sorted = sorted(lst, key=itemgetter('timestamp'))
        #print(lst_sorted)

        x_val = []
        x_val = [x['timestamp'] for x in lst_sorted]
        y_temp = [x['numberOfActivations'] for x in lst_sorted]
        print(x_val)
        # print(y_temp)
      

        # remove invalid number of conn entries
        y_vals = [0 if x is None else x for x in y_temp] # replace none with 0

        # idx = 0
        # trim_lst = []
        # for z in y_vals:
        #     if z == 0:
        #         trim_lst.append(idx)
        #     idx = idx + 1

        # idx = 0
        # for t in trim_lst:
        #     del y_vals[t-idx]
        #     del x_val[t-idx]
        #     idx = idx + 1

     
        #remove invalid dates, before 2018 (1514768400)
        idx = 0
        trim_lst = []
        for z in x_val:
            if z < 1514768400:
                trim_lst.append(idx)
            idx = idx + 1

        idx = 0
        for t in trim_lst:
            #del y_vals[t - idx]
            y_vals.pop(t - idx)
            x_val.pop(t-idx)
            idx = idx + 1

        print(x_val)
        print(y_vals)

        if len(y_vals) > 1:
            plt.clf()
            plt.plot(x_val, y_vals)
            #plt.xticks(np.arange(min(x_val), max(x_val)+1, 1.0))
            # plt.bar(x_val, ydiff, 0.5, color='Red')
            plt.ylabel("Total Number of Resets")
            


            if x_val[len(x_val) - 1] < (time.time() - 604800):  # if last entry more than a week ago sensor is probably dead
                # create plot
                if len(y_vals) <= 3:
                    plt.title("SN " + sensor_serial_number + " DEAD!!" + "\nWarning: 3 or less data points, may not be reliable")
                else: 
                    plt.title("SN " + sensor_serial_number + " DEAD!!")

                
                ax = plt.gca()
            
                plt.gcf().autofmt_xdate()
                ax.xaxis.set_major_locator(mtick.AutoLocator())
                ax.xaxis.set_minor_locator(mtick.AutoMinorLocator())

                ax.xaxis.set_major_formatter(
                mtick.FuncFormatter(lambda pos, _: time.strftime("%m-%d-%Y", time.localtime(pos))))

                plt.tight_layout()
            
                # check whether or not dead_sensors.docx exists and if not then create it
                if os.path.isfile('./dead_sensors.docx') is True:
                    doc = Document('./dead_sensors.docx')
                else:
                    doc = Document()

                # add plot to alive_sensors.docx
                plt.savefig('./matplotlibExample.png')
                doc.add_picture('./matplotlibExample.png', width=Inches(5.5))

                plt.clf()
                plt.cla()
                doc.save('./dead_sensors.docx')

                # add plot to all_sensors.docx after checking if all_sensors.docx exists
                if os.path.isfile('./all_sensors.docx') is True:
                    doc = Document('./all_sensors.docx')
                else:
                    doc = Document()
                doc.add_picture('./matplotlibExample.png', width=Inches(5.5))
                doc.save('./all_sensors.docx')
                
            else:
                # create plot
                if len(y_vals) <= 3:
                    plt.title("SN " + sensor_serial_number + "\nWarning: 3 or less data points, may not be reliable")
                else: 
                    plt.title("SN " + sensor_serial_number)

                doc = Document()
                ax = plt.gca()
            
                plt.gcf().autofmt_xdate()
                ax.xaxis.set_major_locator(mtick.AutoLocator())
                ax.xaxis.set_minor_locator(mtick.AutoMinorLocator())

                ax.xaxis.set_major_formatter(
                mtick.FuncFormatter(lambda pos, _: time.strftime("%m-%d-%Y", time.localtime(pos))))

                plt.tight_layout()

                # check whether or not alive_sensors.docx exists and if not then create it
                if os.path.isfile('./alive_sensors.docx') is True:
                    doc = Document('./alive_sensors.docx')
                else:
                    doc = Document()

                # add plot to alive_sensors.docx
                plt.savefig('./matplotlibExample.png')
                doc.add_picture('./matplotlibExample.png', width=Inches(5.5))

                plt.clf()
                plt.cla()
                doc.save('./alive_sensors.docx')

                # add plot to all_sensors.docx after checking if all_sensors.docx exists
                if os.path.isfile('./all_sensors.docx') is True:
                    doc = Document('./all_sensors.docx')
                else:
                    doc = Document()
                doc.add_picture('./matplotlibExample.png', width=Inches(5.5))
                doc.save('./all_sensors.docx')

if __name__ == "__main__":
    main()
